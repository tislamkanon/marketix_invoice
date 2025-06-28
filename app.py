# === IMPORTS AND SETUP ===
import streamlit as st
import docx
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import pypandoc
import json
import requests
from PIL import Image
import io
import re
import lxml.etree as ET
import tempfile
import zipfile
import shutil

# Set page config as the FIRST Streamlit command
st.set_page_config(page_title="Invoice Generator", page_icon="📄", layout="wide")

# Custom CSS for muted colors, rounded layout, and depth
st.markdown("""
    <style>
    /* General page styling */
    .stApp {
        background-color: #f0f2f6;
        font-family: 'Arial', sans-serif;
    }

    /* Header styling */
    h1 {
        color: #4a5e6a !important;
        text-align: center;
        text-shadow: 1px 1px 2px rgba(0, 0, 0, 0.1);
        padding: 20px;
        background: linear-gradient(145deg, #e6e9ef, #d5d9e0);
        border-radius: 15px;
        box-shadow: 5px 5px 15px rgba(0, 0, 0, 0.1), -5px -5px 15px rgba(255, 255, 255, 0.8);
        margin-bottom: 30px;
    }

    /* Section headers */
    h2 {
        color: #5a7d7c !important;
        margin-top: 20px;
        margin-bottom: 10px;
        padding: 10px 20px;
        background-color: #e2ece9;
        border-radius: 10px;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
    }

    /* Form container styling */
    .stForm {
        background: linear-gradient(145deg, #e6e9ef, #d5d9e0);
        padding: 20px;
        border-radius: 15px;
        box-shadow: 5px 5px 15px rgba(0, 0, 0, 0.1), -5px -5px 15px rgba(255, 255, 255, 0.8);
        margin-bottom: 20px;
    }

    /* Input fields */
    .stTextInput > div > input,
    .stTextArea > div > textarea,
    .stNumberInput > div > input,
    .stDateInput > div > input {
        background-color: #f7f9fc !important;
        border: 1px solid #b0c4c3 !important;
        border-radius: 10px !important;
        padding: 10px !important;
        box-shadow: inset 2px 2px 5px rgba(0, 0, 0, 0.05), inset -2px -2px 5px rgba(255, 255, 255, 0.5) !important;
        color: #4a5e6a !important;
    }

    /* Buttons */
    .stButton > button {
        background: linear-gradient(145deg, #a3bffa, #7f9cfb);
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 10px 20px !important;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
        transition: transform 0.1s ease-in-out;
    }

    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 5px 5px 15px rgba(0, 0, 0, 0.15), -5px -5px 15px rgba(255, 255, 255, 0.9);
    }

    /* Download buttons */
    .stDownloadButton > button {
        background: linear-gradient(145deg, #f4a261, #e76f51);
        color: white !important;
        border-radius: 10px !important;
        padding: 10px 20px !important;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
    }

    /* Checkbox and labels */
    .stCheckbox > label {
        color: #5a7d7c !important;
    }

    /* Error and success messages */
    .stAlert {
        border-radius: 10px !important;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
    }

    /* Items section */
    .stColumn {
        padding: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# === UTILITY CLASSES AND FUNCTIONS ===
class InvoiceData:
    def __init__(self):
        self.client_info = {}
        self.invoice_details = {}
        self.items = []
        self.financials = {}
        self.apply_late_fee = False
        self.mark_as_paid = False
        self.invoice_number = ""
        self.signature = ""

    def to_dict(self):
        return {
            "client_info": self.client_info,
            "invoice_details": self.invoice_details,
            "items": self.items,
            "financials": self.financials,
            "apply_late_fee": self.apply_late_fee,
            "mark_as_paid": self.mark_as_paid,
            "invoice_number": self.invoice_number,
            "signature": self.signature
        }

    @staticmethod
    def from_dict(data):
        invoice = InvoiceData()
        invoice.client_info = data.get("client_info", {})
        invoice.invoice_details = data.get("invoice_details", {})
        invoice.items = data.get("items", [])
        invoice.financials = data.get("financials", {})
        invoice.apply_late_fee = data.get("apply_late_fee", False)
        invoice.mark_as_paid = data.get("mark_as_paid", False)
        invoice.invoice_number = data.get("invoice_number", "")
        invoice.signature = data.get("signature", "")
        return invoice

def format_currency(amount):
    if amount == 0:
        return ""
    elif amount == int(amount):
        return f"Rp {int(amount):,}"
    else:
        return f"Rp {amount:,.2f}"

def get_next_invoice_number():
    count_file = "invoice_count.txt"
    year = "2025"
    if os.path.exists(count_file):
        with open(count_file, 'r') as f:
            try:
                count = int(f.read().strip())
            except ValueError:
                count = 0
    else:
        count = 0
    count += 1
    return f"INV{year}{count:03d}", count

def save_invoice_count(count):
    with open("invoice_count.txt", 'w') as f:
        f.write(str(count))

def validate_date_format(date_str):
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        return True
    except ValueError:
        return False

def save_invoice_data(invoice_data):
    invoice_db = "invoices.json"
    if os.path.exists(invoice_db):
        with open(invoice_db, 'r') as f:
            invoices = json.load(f)
    else:
        invoices = {}

    invoices[invoice_data.invoice_number] = invoice_data.to_dict()
    with open(invoice_db, 'w') as f:
        json.dump(invoices, f, indent=4)

def load_invoice_data():
    invoice_db = "invoices.json"
    if os.path.exists(invoice_db):
        with open(invoice_db, 'r') as f:
            data = json.load(f)
        return {k: InvoiceData.from_dict(v) for k, v in data.items()}
    return {}

def sanitize_filename(name):
    # Remove or replace characters that are invalid in file names
    return re.sub(r'[<>:"/\\|?*]', '_', name).replace(' ', '_')

# === DOCUMENT STYLING FUNCTIONS ===
def set_cell_border(cell, side, color="FFFFFF", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    side_mapping = {
        'top': 'top', 'bottom': 'bottom', 'left': 'left', 'right': 'right'
    }
    border_name = side_mapping.get(side.lower())
    if border_name:
        border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        tcBorders.append(border)

def set_white_borders(cell, sz=4):
    for border in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, border, color="FFFFFF", sz=sz)

def set_cell_font(cell, font_name="Courier New", font_size=10):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def apply_cell_style(cell, bg_color="#ddefd5"):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" />')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    set_white_borders(cell, sz=6)
    set_cell_font(cell)

def style_financial_table(doc, invoice_data):
    financial_table = doc.tables[1]
    for row in financial_table.rows:
        for cell in row.cells:
            set_white_borders(cell)
            set_cell_font(cell)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if invoice_data.apply_late_fee:
        late_fee_cell = financial_table.rows[3].cells[0]
        if "LATE FEE" in late_fee_cell.text:
            original_text = late_fee_cell.text
            late_fee_cell.text = ""
            paragraph = late_fee_cell.paragraphs[0]
            run = paragraph.add_run(original_text)
            run.font.color.rgb = RGBColor.from_string('d95132')
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")

# === INVOICE GENERATION LOGIC ===
def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    return doc

def update_items_table(doc, items):
    items_table = doc.tables[0]
    for i in range(len(items_table.rows)):
        for cell in items_table.rows[i].cells:
            set_white_borders(cell, sz=6)
    while len(items_table.rows) > 2:
        items_table._tbl.remove(items_table.rows[2]._tr)
    placeholder_row = items_table.rows[1]
    for item in items:
        row = items_table.add_row()
        row.cells[0].text = item['description']
        row.cells[1].text = format_currency(item['unit_price'])
        quantity = item['quantity']
        if quantity == int(quantity):
            row.cells[2].text = str(int(quantity))
        else:
            row.cells[2].text = str(quantity)
        row.cells[3].text = format_currency(item['total'])
        for i, cell in enumerate(row.cells):
            apply_cell_style(cell)
            alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, 
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[i]
    items_table._tbl.remove(placeholder_row._tr)
    return doc

def generate_invoice(invoice_data):
    doc = Document('Invoice_Template_MarketixLab.docx')
    replacements = {**invoice_data.client_info, **invoice_data.invoice_details, **invoice_data.financials}
    if invoice_data.apply_late_fee:
        replacements['{{LATE FEE:}}'] = 'LATE FEE'
    else:
        replacements['{{LATE FEE:}}'] = ''
        replacements['[latefee]'] = ''
    doc = replace_placeholders(doc, replacements)
    doc = update_items_table(doc, invoice_data.items)
    style_financial_table(doc, invoice_data)

    if invoice_data.mark_as_paid:
        doc = add_paid_stamp_and_signature(doc)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")
    
    docx_output = io.BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    
    temp_docx = f"temp_{invoice_data.invoice_number}.docx"
    temp_pdf = f"temp_{invoice_data.invoice_number}.pdf"
    doc.save(temp_docx)
    
    pypandoc.convert_file(temp_docx, 'pdf', outputfile=temp_pdf)
    
    pdf_output = io.BytesIO()
    with open(temp_pdf, 'rb') as f:
        pdf_output.write(f.read())
    pdf_output.seek(0)
    
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
    if os.path.exists(temp_pdf):
        os.remove(temp_pdf)
    
    # Generate file names based on paid status and client name
    client_name = sanitize_filename(invoice_data.client_info['{{client_name}}'])
    prefix = "Paid_Invoice" if invoice_data.mark_as_paid else "Invoice"
    base_filename = f"{prefix}_{invoice_data.invoice_number}_{client_name}"
    docx_filename = f"{base_filename}.docx"
    pdf_filename = f"{base_filename}.pdf"
    
    return (docx_output, docx_filename, pdf_output, pdf_filename)

# === STAMP AND SIGNATURE FUNCTIONS ===
# Direct download URLs for the stamp and signature images
PAID_STAMP_URL = "https://drive.google.com/uc?export=download&id=1W9PL0DtP0TUk7IcGiMD_ZuLddtQ8gjNo"
SIGNATURE_URL = "https://drive.google.com/uc?export=download&id=1b6Dcg4spQmvLUMd4neBtLNfdr5l7QtPJ"

def fetch_image(url):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        session = requests.Session()
        response = session.get(url, headers=headers, stream=True, allow_redirects=True)
        
        if response.status_code != 200:
            raise Exception(f"Failed to fetch image from {url}. Status code: {response.status_code}")

        content_type = response.headers.get('Content-Type', '')
        if not content_type.startswith('image/'):
            response_text = response.text
            if "google.com" in response_text and "confirm=" in response_text:
                confirm_match = re.search(r'confirm=([a-zA-Z0-9\-_]+)', response_text)
                if confirm_match:
                    confirm_token = confirm_match.group(1)
                    confirm_url = f"{url}&confirm={confirm_token}"
                    response = session.get(confirm_url, headers=headers, stream=True, allow_redirects=True)
                    content_type = response.headers.get('Content-Type', '')
                    if not content_type.startswith('image/'):
                        response_content = response.text[:200]
                        raise Exception(
                            f"URL {url} still did not return an image after confirmation. "
                            f"Content-Type: {content_type}. "
                            f"Response preview: {response_content}"
                        )
                else:
                    response_content = response_text[:200]
                    raise Exception(
                        f"URL {url} returned a confirmation page, but no confirmation token found. "
                        f"Content-Type: {content_type}. "
                        f"Response preview: {response_content}"
                    )
            else:
                response_content = response_text[:200]
                raise Exception(
                    f"URL {url} did not return an image. "
                    f"Content-Type: {content_type}. "
                    f"Response preview: {response_content}"
                )

        image_data = io.BytesIO(response.content)
        img = Image.open(image_data)
        img.verify()
        image_data.seek(0)
        
        return image_data

    except Exception as e:
        raise Exception(f"Error fetching image from {url}: {str(e)}")

def add_paid_stamp_and_signature(doc):
    try:
        # Fetch images from the URLs
        stamp_data = fetch_image(PAID_STAMP_URL)
        signature_data = fetch_image(SIGNATURE_URL)

        # Save images to temporary files
        stamp_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        signature_tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png')

        # Process stamp image
        stamp_img = Image.open(stamp_data)
        stamp_img.save(stamp_tmp.name, format="PNG")
        stamp_tmp.close()
        if not os.path.exists(stamp_tmp.name):
            raise Exception(f"Stamp temporary file {stamp_tmp.name} does not exist")
        if not os.access(stamp_tmp.name, os.R_OK):
            raise Exception(f"Stamp temporary file {stamp_tmp.name} is not readable")

        # Process signature image
        signature_img = Image.open(signature_data)
        signature_img.save(signature_tmp.name, format="PNG")
        signature_tmp.close()
        if not os.path.exists(signature_tmp.name):
            raise Exception(f"Signature temporary file {signature_tmp.name} does not exist")
        if not os.access(signature_tmp.name, os.R_OK):
            raise Exception(f"Signature temporary file {signature_tmp.name} is not readable")

        # Add the stamp at the end of the document
        stamp_paragraph = doc.add_paragraph()
        stamp_run = stamp_paragraph.add_run()
        stamp_picture = stamp_run.add_picture(stamp_tmp.name, width=Inches(2.17), height=Inches(2.17))

        # Access the run's XML element to find the drawing element
        stamp_run_element = stamp_run._r
        stamp_drawing_elements = stamp_run_element.xpath('.//w:drawing')
        if not stamp_drawing_elements:
            raise Exception("Could not find drawing element for stamp image")
        stamp_drawing = stamp_drawing_elements[0]

        # Find the a:graphic element to preserve the image data
        graphic_elements = stamp_drawing.xpath('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if not graphic_elements:
            raise Exception("Could not find a:graphic element in stamp drawing")
        graphic_xml = ET.tostring(graphic_elements[0], encoding='unicode').replace('\n', '')

        # Use desired positions for stamp
        stamp_horizontal = 5.09 * 914400  # 5.09" in EMUs
        stamp_vertical = 6.64 * 914400    # 6.64" in EMUs

        # Replace the inline drawing with an anchored one using "In Front of Text" wrapping
        stamp_drawing.getparent().replace(stamp_drawing, parse_xml(f"""
            <w:drawing
                xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:posOffset>{int(stamp_horizontal)}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:posOffset>{int(stamp_vertical)}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{int(2.17 * 914400)}" cy="{int(2.17 * 914400)}"/>
                    <wp:effectExtent l="0" t="0" r="0" b="0"/>
                    <wp:wrapTopAndBottom/>
                    <wp:docPr id="1" name="Picture 1"/>
                    <wp:cNvGraphicFramePr/>
                    {graphic_xml}
                </wp:anchor>
            </w:drawing>
        """))

        # Add the signature at the end of the document
        signature_paragraph = doc.add_paragraph()
        signature_run = signature_paragraph.add_run()
        signature_picture = signature_run.add_picture(signature_tmp.name, width=Inches(1.92), height=Inches(1.92))

        # Access the run's XML element to find the drawing element
        signature_run_element = signature_run._r
        signature_drawing_elements = signature_run_element.xpath('.//w:drawing')
        if not signature_drawing_elements:
            raise Exception("Could not find drawing element for signature image")
        signature_drawing = signature_drawing_elements[0]

        # Find the a:graphic element to preserve the image data
        graphic_elements = signature_drawing.xpath('.//a:graphic', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
        if not graphic_elements:
            raise Exception("Could not find a:graphic element in signature drawing")
        graphic_xml = ET.tostring(graphic_elements[0], encoding='unicode').replace('\n', '')

        # Use desired positions for signature
        signature_horizontal = 5.64 * 914400  # 5.64" in EMUs
        signature_vertical = 8.11 * 914400    # 8.11" in EMUs

        # Replace the inline drawing with an anchored one using "In Front of Text" wrapping
        signature_drawing.getparent().replace(signature_drawing, parse_xml(f"""
            <w:drawing
                xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="252" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="page">
                        <wp:posOffset>{int(signature_horizontal)}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="page">
                        <wp:posOffset>{int(signature_vertical)}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{int(1.92 * 914400)}" cy="{int(1.92 * 914400)}"/>
                    <wp:effectExtent l="0" t="0" r="0" b="0"/>
                    <wp:wrapTopAndBottom/>
                    <wp:docPr id="2" name="Picture 2"/>
                    <wp:cNvGraphicFramePr/>
                    {graphic_xml}
                </wp:anchor>
            </w:drawing>
        """))

        # Clean up temporary files
        if os.path.exists(stamp_tmp.name):
            os.remove(stamp_tmp.name)
        if os.path.exists(signature_tmp.name):
            os.remove(signature_tmp.name)

        return doc

    except Exception as e:
        # Clean up temporary files in case of failure
        if 'stamp_tmp' in locals() and os.path.exists(stamp_tmp.name):
            os.remove(stamp_tmp.name)
        if 'signature_tmp' in locals() and os.path.exists(signature_tmp.name):
            os.remove(signature_tmp.name)
        raise Exception(f"Failed to add stamp and signature: {str(e)}")

# === STREAMLIT UI AND APP LOGIC ===
st.title("📄 Invoice Generator")
st.markdown("Create professional invoices with ease using this streamlined tool.")

tab1, tab2 = st.tabs(["Create Invoice", "View Invoices"])

with tab1:
    if 'item_list' not in st.session_state:
        st.session_state.item_list = [{"description": "", "unit_price": 0.0, "quantity": 0.0}]

    if 'use_today' not in st.session_state:
        st.session_state.use_today = True

    if 'manual_invoice_date' not in st.session_state:
        st.session_state.manual_invoice_date = datetime.now()

    st.header("Client Information")
    with st.form(key="client_form"):
        client_name = st.text_input("Client Name", placeholder="Enter client name")
        client_phone = st.text_input("Client Phone", placeholder="Enter phone number")
        client_email = st.text_input("Client Email", placeholder="Enter email")
        client_address = st.text_area("Client Address", placeholder="Enter address")
        client_submit = st.form_submit_button("Save Client Info")

    st.header("Invoice Details")
    with st.form(key="invoice_form"):
        default_invoice_number, invoice_count = get_next_invoice_number()
        invoice_number = st.text_input("Invoice Number", value=default_invoice_number, help="Invoice number must start with 'INV2025'")
        
        st.session_state.use_today = st.checkbox("Use Today's Date", value=st.session_state.use_today, key="use_today_checkbox")
        if st.session_state.use_today:
            invoice_date = datetime.now().strftime("%d.%m.%Y")
            st.write(f"Invoice Date: {invoice_date}")
        else:
            st.session_state.manual_invoice_date = st.date_input(
                "Select Invoice Date",
                value=st.session_state.manual_invoice_date,
                key="manual_invoice_date_picker"
            )
            invoice_date = st.session_state.manual_invoice_date.strftime("%d.%m.%Y")
            st.write(f"Selected Invoice Date: {invoice_date}")
        
        due_date_obj = st.date_input("Select Due Date", value=datetime.now(), key="due_date_picker")
        due_date = due_date_obj.strftime("%d.%m.%Y")
        
        invoice_submit = st.form_submit_button("Save Invoice Details")

    st.header("Items")
    if not isinstance(st.session_state.item_list, list):
        st.warning("Item list was corrupted. Resetting to default.")
        st.session_state.item_list = [{"description": "", "unit_price": 0.0, "quantity": 0.0}]

    def add_item():
        st.session_state.item_list.append({"description": "", "unit_price": 0.0, "quantity": 0.0})

    def remove_item(index):
        if len(st.session_state.item_list) > 1:
            st.session_state.item_list.pop(index)

    for i in range(len(st.session_state.item_list)):
        col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
        with col1:
            st.session_state.item_list[i]["description"] = st.text_input(
                f"Description {i+1}",
                value=st.session_state.item_list[i]["description"],
                key=f"desc_{i}"
            )
        with col2:
            st.session_state.item_list[i]["unit_price"] = st.number_input(
                f"Unit Price {i+1}",
                min_value=0.0,
                value=st.session_state.item_list[i]["unit_price"],
                key=f"price_{i}"
            )
        with col3:
            st.session_state.item_list[i]["quantity"] = st.number_input(
                f"Quantity {i+1}",
                min_value=0.0,
                value=st.session_state.item_list[i]["quantity"],
                key=f"qty_{i}"
            )
        with col4:
            if st.button("✕", key=f"delete_{i}"):
                remove_item(i)

    st.button("Add Item", on_click=add_item)

    st.header("Financial Details")
    with st.form(key="financial_form"):
        tax_rate = st.number_input("Tax Rate (%)", min_value=0.0, value=0.0, help="Enter tax rate as a percentage")
        discount = st.number_input("Discount Amount", min_value=0.0, value=0.0, help="Enter discount amount in Rp")
        apply_late_fee = st.checkbox("Apply Late Fee (2%)", value=False, help="Check to apply a 2% late fee")
        financial_submit = st.form_submit_button("Save Financial Details")

    if st.button("Generate Invoice"):
        try:
            if not all([client_name, client_phone, client_email, client_address]):
                st.error("All client info fields are required")
            elif not all([invoice_number, invoice_date, due_date]):
                st.error("All invoice details are required")
            elif not invoice_number.startswith("INV2025"):
                st.error("Invoice number must start with 'INV2025'")
            elif not validate_date_format(invoice_date):
                st.error("Invoice date must be in the format dd.mm.yyyy (e.g., 21.04.2025)")
            elif not validate_date_format(due_date):
                st.error("Due date must be in the format dd.mm.yyyy (e.g., 28.04.2025)")
            elif not st.session_state.item_list or not any(item["description"] and item["unit_price"] > 0 and item["quantity"] > 0 for item in st.session_state.item_list):
                st.error("At least one valid item is required")
            else:
                invoice_data = InvoiceData()
                invoice_data.client_info = {
                    '{{client_name}}': client_name,
                    '{{client_phone}}': client_phone,
                    '{{client_email}}': client_email,
                    '{{client_address}}': client_address
                }
                invoice_data.invoice_details = {
                    '{{invoice_number}}': invoice_number,
                    '{{invoice_date}}': invoice_date,
                    '{{due_date}}': due_date
                }
                invoice_data.items = [
                    {
                        'description': item['description'],
                        'unit_price': item['unit_price'],
                        'quantity': item['quantity'],
                        'total': item['unit_price'] * item['quantity']
                    } for item in st.session_state.item_list if item['description'] and item['unit_price'] > 0 and item['quantity'] > 0
                ]
                subtotal = sum(item['total'] for item in invoice_data.items)
                tax = subtotal * (tax_rate / 100)
                invoice_data.apply_late_fee = apply_late_fee
                late_fee = subtotal * 0.02 if apply_late_fee else 0
                total = subtotal + tax - discount + late_fee
                invoice_data.financials = {
                    '[subtotal]': format_currency(subtotal),
                    '[tax]': format_currency(tax),
                    '[discount]': format_currency(discount),
                    '[latefee]': format_currency(late_fee),
                    '[grandtotal]': format_currency(total)
                }
                invoice_data.invoice_number = invoice_number
                save_invoice_data(invoice_data)
                docx_output, docx_filename, pdf_output, pdf_filename = generate_invoice(invoice_data)
                if invoice_number == default_invoice_number:
                    save_invoice_count(invoice_count)
                st.success(f"Invoice {invoice_number} generated and saved successfully!")
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        label="Download Invoice (DOCX)",
                        data=docx_output,
                        file_name=docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        label="Download Invoice (PDF)",
                        data=pdf_output,
                        file_name=pdf_filename,
                        mime="application/pdf"
                    )
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

with tab2:
    st.header("Previously Generated Invoices")
    invoices = load_invoice_data()
    if not invoices:
        st.info("No invoices found.")
    else:
        # Add filter for paid/unpaid invoices
        filter_option = st.selectbox(
            "Filter Invoices",
            ["All Invoices", "Paid Invoices", "Unpaid Invoices"],
            key="invoice_filter"
        )

        # Filter invoices based on selection
        if filter_option == "Paid Invoices":
            filtered_invoices = {k: v for k, v in invoices.items() if v.mark_as_paid}
        elif filter_option == "Unpaid Invoices":
            filtered_invoices = {k: v for k, v in invoices.items() if not v.mark_as_paid}
        else:
            filtered_invoices = invoices

        if not filtered_invoices:
            st.info(f"No {filter_option.lower()} found.")
        else:
            invoice_numbers = list(filtered_invoices.keys())
            selected_invoice = st.selectbox("Select an Invoice", invoice_numbers, key="select_invoice")
            if selected_invoice:
                invoice_data = filtered_invoices[selected_invoice]
                st.write(f"**Invoice Number:** {invoice_data.invoice_number}")
                st.write(f"**Client Name:** {invoice_data.client_info['{{client_name}}']}")
                st.write(f"**Date:** {invoice_data.invoice_details['{{invoice_date}}']}")
                st.write(f"**Due Date:** {invoice_data.invoice_details['{{due_date}}']}")
                st.write(f"**Total:** {invoice_data.financials['[grandtotal]']}")
                st.write(f"**Paid Status:** {'Paid' if invoice_data.mark_as_paid else 'Not Paid'}")

                if not invoice_data.mark_as_paid:
                    if st.button(f"Mark {selected_invoice} as Paid"):
                        invoice_data.mark_as_paid = True
                        save_invoice_data(invoice_data)
                        st.success(f"Invoice {selected_invoice} marked as paid!")
                        st.experimental_rerun()

                if st.button(f"Download {selected_invoice}"):
                    try:
                        docx_output, docx_filename, pdf_output, pdf_filename = generate_invoice(invoice_data)
                        st.success(f"Invoice {selected_invoice} generated successfully!")
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                label="Download Invoice (DOCX)",
                                data=docx_output,
                                file_name=docx_filename,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        with col2:
                            st.download_button(
                                label="Download Invoice (PDF)",
                                data=pdf_output,
                                file_name=pdf_filename,
                                mime="application/pdf"
                            )
                    except Exception as e:
                        st.error(f"An error occurred: {str(e)}")
