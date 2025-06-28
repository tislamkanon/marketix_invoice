import docx
from docx import Document
from datetime import datetime
import uuid
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
from tkcalendar import DateEntry
import os

class InvoiceData:
    def __init__(self):
        self.client_info = {}
        self.invoice_details = {}
        self.items = []
        self.financials = {}
        self.apply_late_fee = False
        self.invoice_number = ""

def format_currency(amount):
    if amount == 0:
        return ""
    elif amount == int(amount):
        return f"Rp {int(amount):,}"
    else:
        return f"Rp {amount:,.2f}"

def set_cell_border(cell, side, color="FFFFFF", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    side_mapping = {
        'top': 'top', 'bottom': 'bottom', 'left': 'left', 'right': 'right'
    }
    border_name = side_mapping.get(side.lower())
    if border_name:
        border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        tcBorders.append(border)

def set_white_borders(cell, sz=4):
    for border in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, border, color="FFFFFF", sz=sz)

def set_cell_font(cell, font_name="Courier New", font_size=10):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def apply_cell_style(cell, bg_color="#ddefd5"):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" />')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    set_white_borders(cell, sz=6)
    set_cell_font(cell)

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    return doc

def update_items_table(doc, items):
    items_table = doc.tables[0]
    for i in range(len(items_table.rows)):
        for cell in items_table.rows[i].cells:
            set_white_borders(cell, sz=6)
    while len(items_table.rows) > 2:
        items_table._tbl.remove(items_table.rows[2]._tr)
    placeholder_row = items_table.rows[1]
    for item in items:
        row = items_table.add_row()
        row.cells[0].text = item['description']
        row.cells[1].text = format_currency(item['unit_price'])
        quantity = item['quantity']
        if quantity == int(quantity):
            row.cells[2].text = str(int(quantity))
        else:
            row.cells[2].text = str(quantity)
        row.cells[3].text = format_currency(item['total'])
        for i, cell in enumerate(row.cells):
            apply_cell_style(cell)
            alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT,
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[i]
    items_table._tbl.remove(placeholder_row._tr)
    return doc

def style_financial_table(doc, invoice_data):
    financial_table = doc.tables[1]
    for row in financial_table.rows:
        for cell in row.cells:
            set_white_borders(cell)
            set_cell_font(cell)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if invoice_data.apply_late_fee:
        late_fee_cell = financial_table.rows[3].cells[0]
        if "LATE FEE" in late_fee_cell.text:
            original_text = late_fee_cell.text
            late_fee_cell.text = ""
            paragraph = late_fee_cell.paragraphs[0]
            run = paragraph.add_run(original_text)
            run.font.color.rgb = RGBColor.from_string('d95132')
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")

def generate_invoice(invoice_data):
    doc = Document('Invoice_Template_MarketixLab.docx')
    replacements = {**invoice_data.client_info, **invoice_data.invoice_details, **invoice_data.financials}
    if invoice_data.apply_late_fee:
        replacements['{{LATE FEE:}}'] = 'LATE FEE'
    else:
        replacements['{{LATE FEE:}}'] = ''
        replacements['[latefee]'] = ''
    doc = replace_placeholders(doc, replacements)
    doc = update_items_table(doc, invoice_data.items)
    style_financial_table(doc, invoice_data)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")
    output_filename = f"Invoice_{invoice_data.invoice_number}.docx"
    doc.save(output_filename)
    messagebox.showinfo("Success", f"Invoice saved as {output_filename}")

def get_next_invoice_number():
    count_file = "invoice_count.txt"
    year = "2025"
    if os.path.exists(count_file):
        with open(count_file, 'r') as f:
            try:
                count = int(f.read().strip())
            except ValueError:
                count = 0
    else:
        count = 0
    count += 1
    return f"INV{year}{count:03d}", count

def save_invoice_count(count):
    with open("invoice_count.txt", 'w') as f:
        f.write(str(count))

class InvoiceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Invoice Generator")
        self.root.geometry("700x900")
        self.root.configure(bg="#e5e7eb")

        self.invoice_data = InvoiceData()
        self.item_rows = []

        style = ttk.Style()
        style.configure('TFrame', background='#e5e7eb')
        style.configure('TLabel', background='#e5e7eb', font=('Arial', 11, 'bold'), foreground='#4b5563')
        style.configure('TEntry', font=('Arial', 11))
        style.configure('TButton', font=('Arial', 11, 'bold'), foreground='#ffffff')
        style.configure('Custom.TButton', background='#6b7280', foreground='#ffffff')
        style.map('Custom.TButton', background=[('active', '#4b5563')])

        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill='both', expand=True)

        canvas = tk.Canvas(main_frame, bg="#e5e7eb", highlightthickness=0)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)

        self.scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.bind_all("<MouseWheel>", lambda event: canvas.yview_scroll(int(-1 * (event.delta / 120)), "units"))

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        header_frame = ttk.Frame(self.scrollable_frame)
        header_frame.pack(fill='x', pady=(20, 10))
        header_label = tk.Label(header_frame, text="📄 Invoice Generator", font=('Arial', 24, 'bold'),
                               bg="#e5e7eb", fg="#4b5563")
        header_label.pack()

        self.client_section = ttk.Frame(self.scrollable_frame)
        self.client_section.pack(fill='x', padx=20, pady=10)
        self.create_section_header(self.client_section, "Client Information")
        self.setup_client_section()

        self.invoice_section = ttk.Frame(self.scrollable_frame)
        self.invoice_section.pack(fill='x', padx=20, pady=10)
        self.create_section_header(self.invoice_section, "Invoice Details")
        self.setup_invoice_section()

        self.items_section = ttk.Frame(self.scrollable_frame)
        self.items_section.pack(fill='x', padx=20, pady=10)
        self.create_section_header(self.items_section, "Items")
        self.setup_items_section()

        self.financial_section = ttk.Frame(self.scrollable_frame)
        self.financial_section.pack(fill='x', padx=20, pady=10)
        self.create_section_header(self.financial_section, "Financial Details")
        self.setup_financial_section()

        button_frame = ttk.Frame(self.scrollable_frame)
        button_frame.pack(pady=30)
        self.generate_button = ttk.Button(button_frame, text="Generate Invoice", style='Custom.TButton',
                                         command=self.collect_data_and_generate)
        self.generate_button.pack()

    def create_section_header(self, frame, text):
        header_frame = ttk.Frame(frame)
        header_frame.pack(fill='x', pady=(10, 5))
        header_frame.configure(style='TFrame')
        canvas = tk.Canvas(header_frame, height=40, bg="#e5e7eb", highlightthickness=0)
        canvas.pack(fill='x', expand=True)
        canvas.create_rectangle(10, 10, 680, 40, fill="#d1d5db", outline="#d1d5db")
        canvas.create_text(20, 25, text=text, anchor='w', font=('Arial', 16, 'bold'), fill="#4b5563")

    def setup_client_section(self):
        frame = ttk.Frame(self.client_section)
        frame.pack(fill='x', padx=10, pady=10)
        fields = ["Client Name:", "Client Phone:", "Client Email:", "Client Address:"]
        self.client_entries = {}
        for i, field in enumerate(fields):
            ttk.Label(frame, text=field).grid(row=i, column=0, sticky='w', pady=5, padx=5)
            entry = ttk.Entry(frame, width=50)
            entry.grid(row=i, column=1, sticky='ew', pady=5, padx=5)
            self.client_entries[field] = entry
        frame.columnconfigure(1, weight=1)

    def setup_invoice_section(self):
        frame = ttk.Frame(self.invoice_section)
        frame.pack(fill='x', padx=10, pady=10)
        ttk.Label(frame, text="Invoice Number:").grid(row=0, column=0, sticky='w', pady=5, padx=5)
        self.invoice_number = ttk.Entry(frame, width=30)
        self.invoice_number.grid(row=0, column=1, sticky='w', pady=5, padx=5)
        next_invoice_number, _ = get_next_invoice_number()
        self.invoice_number.insert(0, next_invoice_number)

        ttk.Label(frame, text="Invoice Date:").grid(row=1, column=0, sticky='w', pady=5, padx=5)
        self.use_today = tk.IntVar(value=1)
        ttk.Checkbutton(frame, text="Use Today's Date", variable=self.use_today,
                        command=self.toggle_date_entry).grid(row=1, column=1, sticky='w', pady=5, padx=5)

        self.date_frame = ttk.Frame(frame)
        self.date_frame.grid(row=2, column=1, sticky='w', pady=5, padx=5)
        self.invoice_date = DateEntry(self.date_frame, width=20, date_pattern='dd.mm.yyyy', state='disabled')
        self.invoice_date.pack(side='left')

        ttk.Label(frame, text="Due Date:").grid(row=3, column=0, sticky='w', pady=5, padx=5)
        self.due_date = DateEntry(frame, width=20, date_pattern='dd.mm.yyyy')
        self.due_date.grid(row=3, column=1, sticky='w', pady=5, padx=5)

    def toggle_date_entry(self):
        if self.use_today.get() == 1:
            self.invoice_date.config(state='disabled')
        else:
            self.invoice_date.config(state='normal')

    def setup_items_section(self):
        frame = ttk.Frame(self.items_section)
        frame.pack(fill='x', padx=10, pady=10)
        ttk.Label(frame, text="Description", width=30).grid(row=0, column=0, padx=5, pady=5)
        ttk.Label(frame, text="Unit Price", width=15).grid(row=0, column=1, padx=5, pady=5)
        ttk.Label(frame, text="Quantity", width=10).grid(row=0, column=2, padx=5, pady=5)

        items_frame = ttk.Frame(frame)
        items_frame.grid(row=1, column=0, columnspan=4, sticky='nsew', pady=5)
        self.items_canvas = tk.Canvas(items_frame, bg="#e5e7eb", highlightthickness=0)
        self.items_canvas.grid(row=0, column=0, sticky='nsew')
        scrollbar = ttk.Scrollbar(items_frame, orient="vertical", command=self.items_canvas.yview)
        scrollbar.grid(row=0, column=1, sticky='ns')
        self.items_canvas.configure(yscrollcommand=scrollbar.set)
        items_frame.columnconfigure(0, weight=1)

        self.items_subframe = ttk.Frame(self.items_canvas)
        self.items_canvas.create_window((0, 0), window=self.items_subframe, anchor='nw')
        self.items_subframe.bind("<Configure>", lambda e: self.items_canvas.configure(
            scrollregion=self.items_canvas.bbox("all")))

        self.add_item_row()
        ttk.Button(frame, text="Add Item", style='Custom.TButton', command=self.add_item_row).grid(row=2, column=0, pady=10, sticky='w')

    def add_item_row(self):
        row_index = len(self.item_rows)
        frame = ttk.Frame(self.items_subframe)
        frame.pack(fill='x', pady=2)
        description = ttk.Entry(frame, width=30)
        description.pack(side='left', padx=5)
        price = ttk.Entry(frame, width=15)
        price.pack(side='left', padx=5)
        quantity = ttk.Entry(frame, width=10)
        quantity.pack(side='left', padx=5)
        delete_btn = ttk.Button(frame, text="✕", width=3, style='Custom.TButton',
                               command=lambda f=frame, idx=row_index: self.delete_item_row(f, idx))
        delete_btn.pack(side='left', padx=5)
        self.item_rows.append({
            'frame': frame,
            'description': description,
            'price': price,
            'quantity': quantity,
            'delete_btn': delete_btn
        })

    def delete_item_row(self, frame, index):
        if len(self.item_rows) > 1:
            frame.destroy()
            self.item_rows.pop(index)
            for i, row in enumerate(self.item_rows):
                row['delete_btn'].configure(command=lambda f=row['frame'], idx=i: self.delete_item_row(f, idx))

    def setup_financial_section(self):
        frame = ttk.Frame(self.financial_section)
        frame.pack(fill='x', padx=10, pady=10)
        ttk.Label(frame, text="Tax Rate (%):").grid(row=0, column=0, sticky='w', pady=5, padx=5)
        self.tax_rate = ttk.Entry(frame, width=20)
        self.tax_rate.grid(row=0, column=1, sticky='w', pady=5, padx=5)
        self.tax_rate.insert(0, "0")

        ttk.Label(frame, text="Discount Amount:").grid(row=1, column=0, sticky='w', pady=5, padx=5)
        self.discount = ttk.Entry(frame, width=20)
        self.discount.grid(row=1, column=1, sticky='w', pady=5, padx=5)
        self.discount.insert(0, "0")

        ttk.Label(frame, text="Apply Late Fee (2%):").grid(row=2, column=0, sticky='w', pady=5, padx=5)
        self.late_fee_var = tk.IntVar(value=0)
        ttk.Checkbutton(frame, text="Yes", variable=self.late_fee_var).grid(row=2, column=1, sticky='w', pady=5, padx=5)

    def collect_data_and_generate(self):
        try:
            self.invoice_data.client_info = {
                '{{client_name}}': self.client_entries["Client Name:"].get(),
                '{{client_phone}}': self.client_entries["Client Phone:"].get(),
                '{{client_email}}': self.client_entries["Client Email:"].get(),
                '{{client_address}}': self.client_entries["Client Address:"].get()
            }
            if not all(self.invoice_data.client_info.values()):
                messagebox.showerror("Error", "All client info fields are required")
                return
            self.invoice_data.invoice_number = self.invoice_number.get()
            if not self.invoice_data.invoice_number.startswith("INV2025"):
                messagebox.showerror("Error", "Invoice number must start with 'INV2025'")
                return
            if self.use_today.get() == 1:
                invoice_date = datetime.now().strftime("%d.%m.%Y")
            else:
                invoice_date = self.invoice_date.get()
            self.invoice_data.invoice_details = {
                '{{invoice_number}}': self.invoice_data.invoice_number,
                '{{invoice_date}}': invoice_date,
                '{{due_date}}': self.due_date.get()
            }
            self.invoice_data.items = []
            for row in self.item_rows:
                description = row['description'].get()
                try:
                    price = float(row['price'].get().replace(',', ''))
                    quantity = float(row['quantity'].get().replace(',', ''))
                except ValueError:
                    messagebox.showerror("Error", "Price and quantity must be numbers")
                    return
                if description and price > 0 and quantity > 0:
                    self.invoice_data.items.append({
                        'description': description,
                        'unit_price': price,
                        'quantity': quantity,
                        'total': price * quantity
                    })
            if not self.invoice_data.items:
                messagebox.showerror("Error", "At least one item is required")
                return
            subtotal = sum(item['total'] for item in self.invoice_data.items)
            try:
                tax_rate = float(self.tax_rate.get()) / 100
                discount = float(self.discount.get().replace(',', ''))
            except ValueError:
                messagebox.showerror("Error", "Tax rate and discount must be numbers")
                return
            tax = subtotal * tax_rate
            self.invoice_data.apply_late_fee = bool(self.late_fee_var.get())
            late_fee = subtotal * 0.02 if self.invoice_data.apply_late_fee else 0
            total = subtotal + tax - discount + late_fee
            self.invoice_data.financials = {
                '[subtotal]': format_currency(subtotal),
                '[tax]': format_currency(tax),
                '[discount]': format_currency(discount),
                '[latefee]': format_currency(late_fee),
                '[grandtotal]': format_currency(total)
            }
            generate_invoice(self.invoice_data)
            _, count = get_next_invoice_number()
            save_invoice_count(count)
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

def main():
    root = tk.Tk()
    app = InvoiceApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
